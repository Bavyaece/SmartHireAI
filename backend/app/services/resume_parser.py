import io
import logging
import re
from docx import Document

logger = logging.getLogger(__name__)

_OCR_ENGINE = None


def _clean_chunk(text: str | None) -> str:
    if not text:
        return ""
    text = text.replace("\x00", " ").replace("\xa0", " ")
    text = text.replace("\ufb01", "fi").replace("\ufb02", "fl")
    return text


def _word_score(text: str) -> int:
    return len(re.findall(r"[A-Za-z]{2,}", text or ""))


def _get_ocr_engine():
    """Lazy-load RapidOCR once (heavy models)."""
    global _OCR_ENGINE
    if _OCR_ENGINE is not None:
        return _OCR_ENGINE
    try:
        from rapidocr_onnxruntime import RapidOCR
        import numpy as np
        from PIL import Image

        _OCR_ENGINE = ("rapid", RapidOCR(), np, Image)
        logger.info("RapidOCR engine loaded")
        return _OCR_ENGINE
    except Exception as e:
        logger.warning("RapidOCR unavailable: %s", e)
        try:
            import pytesseract
            from PIL import Image

            _OCR_ENGINE = ("tesseract", pytesseract, None, Image)
            return _OCR_ENGINE
        except Exception as e2:
            logger.warning("No OCR engine available: %s", e2)
            _OCR_ENGINE = False
            return None


def _ocr_pdf_with_pymupdf(content: bytes, max_pages: int = 5) -> str:
    """OCR scanned PDFs by rendering each page to an image."""
    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz
        except ImportError:
            return ""

    engine = _get_ocr_engine()
    if not engine:
        return ""

    doc = fitz.open(stream=content, filetype="pdf")
    page_texts: list[str] = []

    try:
        for i, page in enumerate(doc):
            if i >= max_pages:
                break
            # ~216 DPI for clearer OCR on resume scans
            pix = page.get_pixmap(matrix=fitz.Matrix(2.5, 2.5), alpha=False)
            png_bytes = pix.tobytes("png")

            kind = engine[0]
            if kind == "rapid":
                _, ocr, np, Image = engine
                img = Image.open(io.BytesIO(png_bytes)).convert("RGB")
                result, _ = ocr(np.array(img))
                if result:
                    lines = [line[1] for line in result if len(line) > 1 and line[1]]
                    page_texts.append("\n".join(lines))
            else:
                _, pytesseract, _, Image = engine
                img = Image.open(io.BytesIO(png_bytes))
                page_texts.append(pytesseract.image_to_string(img) or "")
    except Exception as e:
        logger.exception("OCR failed: %s", e)
        return ""
    finally:
        doc.close()

    return _clean_chunk("\n".join(page_texts))


def _native_pdf_text(content: bytes) -> str:
    candidates: list[str] = []

    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz
        except ImportError:
            fitz = None

    if fitz is not None:
        try:
            doc = fitz.open(stream=content, filetype="pdf")
            parts = []
            for page in doc:
                t = page.get_text("text") or ""
                if _word_score(t) < 8:
                    blocks = page.get_text("blocks")
                    if isinstance(blocks, list):
                        t = "\n".join(
                            b[4] for b in blocks if len(b) > 4 and isinstance(b[4], str)
                        )
                parts.append(_clean_chunk(t))
            doc.close()
            candidates.append("\n".join(parts))
        except Exception:
            pass

    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        parts = []
        for page in reader.pages:
            try:
                text = page.extract_text() or ""
            except Exception:
                text = ""
            parts.append(_clean_chunk(text))
        candidates.append("\n".join(parts))
    except Exception:
        pass

    try:
        from pdfminer.high_level import extract_text as pdfminer_extract

        candidates.append(_clean_chunk(pdfminer_extract(io.BytesIO(content)) or ""))
    except Exception:
        pass

    return max(candidates, key=_word_score) if candidates else ""


def _pdf_has_images(content: bytes) -> bool:
    try:
        import pymupdf as fitz
    except ImportError:
        try:
            import fitz
        except ImportError:
            return False
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        has = any(page.get_images() for page in doc)
        doc.close()
        return has
    except Exception:
        return False


def extract_text_from_pdf(content: bytes) -> str:
    """
    PDF-only pipeline:
    1) Native text extraction
    2) Automatic OCR if text is weak / PDF looks scanned
    """
    native = _native_pdf_text(content)
    native_score = _word_score(native)

    # Always OCR when little text OR page is image-heavy with weak text
    needs_ocr = native_score < 40 or (native_score < 80 and _pdf_has_images(content))

    if needs_ocr:
        logger.info("Running OCR on PDF (native words=%s)", native_score)
        ocr_text = _ocr_pdf_with_pymupdf(content)
        ocr_score = _word_score(ocr_text)
        logger.info("OCR finished (words=%s)", ocr_score)
        if ocr_score > native_score:
            return ocr_text

    return native or ""


def extract_text_from_docx(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text and p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text_from_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except UnicodeDecodeError:
            continue
    return content.decode("utf-8", errors="ignore")


def extract_text(filename: str, content: bytes) -> str:
    if not content:
        raise ValueError("Empty file uploaded.")

    lower = (filename or "").lower().strip()

    if lower.endswith(".pdf") or content[:4] == b"%PDF":
        return extract_text_from_pdf(content)
    if lower.endswith(".docx") or content[:2] == b"PK":
        return extract_text_from_docx(content)
    if lower.endswith(".txt"):
        return extract_text_from_txt(content)
    if lower.endswith(".doc"):
        raise ValueError("Please upload a PDF resume.")

    if content[:4] == b"%PDF":
        return extract_text_from_pdf(content)

    raise ValueError("Please upload a PDF resume.")


def normalize_text(text: str) -> str:
    text = _clean_chunk(text)
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", text)
    text = text.replace("\n", " ")
    text = re.sub(r"\s+", " ", text)
    return text.strip()


def extraction_quality(text: str) -> dict:
    return {
        "chars": len(text),
        "letters": len(re.findall(r"[A-Za-z]", text)),
        "words": len(re.findall(r"[A-Za-z]{2,}", text)),
    }
